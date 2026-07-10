import io

import docx
import pytest
from pdf_fixture import build_minimal_pdf
from PIL import Image

from tmis.document_intelligence.ingestion.docx_parser import DocxParser
from tmis.document_intelligence.ingestion.eml_parser import EmlParser
from tmis.document_intelligence.ingestion.exceptions import (
    DocumentValidationError,
    UnsupportedContentTypeError,
    VirusDetectedError,
)
from tmis.document_intelligence.ingestion.image_parser import ImageParser
from tmis.document_intelligence.ingestion.pdf_parser import PdfParser
from tmis.document_intelligence.ingestion.registry import IngestionRegistry
from tmis.document_intelligence.ingestion.txt_parser import TxtParser
from tmis.document_intelligence.ingestion.validation import DocumentValidator
from tmis.document_intelligence.ingestion.virus_scan import NullVirusScanner
from tmis.document_intelligence.schemas.document import IngestedDocument


def _minimal_png_bytes() -> bytes:
    image = Image.new("RGB", (10, 5), color="white")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _minimal_docx_bytes() -> bytes:
    document = docx.Document()
    document.add_paragraph("Le contrat de bail commercial.")
    document.add_paragraph("Fait à Paris.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestDocumentValidator:
    def test_rejects_empty_file(self) -> None:
        with pytest.raises(DocumentValidationError):
            DocumentValidator().validate("empty.txt", b"")

    def test_rejects_oversized_file(self) -> None:
        with pytest.raises(DocumentValidationError):
            DocumentValidator(max_size_bytes=10).validate("big.txt", b"x" * 11)

    def test_accepts_valid_file(self) -> None:
        DocumentValidator().validate("ok.txt", b"hello")


class TestNullVirusScanner:
    def test_never_raises(self) -> None:
        NullVirusScanner().scan("anything.txt", b"content")


def test_virus_detected_error_message_includes_filename_and_reason() -> None:
    error = VirusDetectedError("bad.txt", "eicar signature")
    assert "bad.txt" in str(error)
    assert "eicar signature" in str(error)


class TestTxtParser:
    def test_parses_utf8_text(self) -> None:
        doc = TxtParser().parse("id-1", "note.txt", "Résiliation du bail.".encode())
        assert doc.text == "Résiliation du bail."
        assert doc.content_type == "text/plain"
        assert doc.page_count == 1

    def test_supports_only_plain_text(self) -> None:
        parser = TxtParser()
        assert parser.supports("text/plain")
        assert not parser.supports("application/pdf")


class TestPdfParser:
    def test_extracts_text_from_minimal_pdf(self) -> None:
        pdf_bytes = build_minimal_pdf("Hello World")
        doc = PdfParser().parse("id-1", "doc.pdf", pdf_bytes)
        assert "Hello World" in doc.text
        assert doc.page_count == 1
        assert doc.content_type == "application/pdf"


class TestDocxParser:
    def test_extracts_paragraphs(self) -> None:
        docx_bytes = _minimal_docx_bytes()
        doc = DocxParser().parse("id-1", "doc.docx", docx_bytes)
        assert "Le contrat de bail commercial." in doc.text
        assert "Fait à Paris." in doc.text


class TestImageParser:
    def test_captures_dimensions_and_leaves_text_empty(self) -> None:
        png_bytes = _minimal_png_bytes()
        doc = ImageParser().parse("id-1", "scan.png", png_bytes)
        assert doc.text == ""
        assert doc.metadata["width"] == "10"
        assert doc.metadata["height"] == "5"


class TestEmlParser:
    def test_is_prepared_but_not_implemented(self) -> None:
        parser = EmlParser()
        assert parser.supports("message/rfc822")
        with pytest.raises(NotImplementedError):
            parser.parse("id-1", "mail.eml", b"From: a@b.com")


class TestIngestionRegistry:
    def test_dispatches_to_matching_parser(self) -> None:
        registry = IngestionRegistry()
        doc = registry.parse("id-1", "note.txt", "text/plain", b"hello")
        assert doc.text == "hello"

    def test_raises_for_unsupported_content_type(self) -> None:
        registry = IngestionRegistry()
        with pytest.raises(UnsupportedContentTypeError):
            registry.parse("id-1", "file.xyz", "application/x-unknown", b"data")

    def test_register_adds_a_custom_parser(self) -> None:
        class _CustomParser:
            content_types: tuple[str, ...] = ("application/x-custom",)

            def supports(self, content_type: str) -> bool:
                return content_type in self.content_types

            def parse(
                self, document_id: str, filename: str, raw_bytes: bytes
            ) -> IngestedDocument:
                return IngestedDocument(
                    id=document_id,
                    filename=filename,
                    content_type="application/x-custom",
                    text="custom",
                    page_count=1,
                )

        registry = IngestionRegistry()
        registry.register(_CustomParser())
        doc = registry.parse("id-1", "f.custom", "application/x-custom", b"x")
        assert doc.text == "custom"
