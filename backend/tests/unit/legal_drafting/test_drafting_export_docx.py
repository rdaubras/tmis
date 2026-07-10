import io

import docx

from tmis.legal_drafting.citations.schemas import DraftCitation
from tmis.legal_drafting.documents.schemas import Document
from tmis.legal_drafting.export.docx_exporter import DocxExporter
from tmis.legal_drafting.export.schemas import ExportFormat
from tmis.legal_drafting.paragraphs.schemas import Paragraph
from tmis.legal_drafting.sections.schemas import Section
from tmis.legal_drafting.templates.schemas import DocumentType


def _document() -> Document:
    paragraph = Paragraph(
        id="p1", section_key="facts", order=0, text="Le contrat a été rompu.", origin="kernel",
    )
    section = Section(id="s1", key="facts", title="Faits", order=0, paragraphs=[paragraph])
    citation = DraftCitation(
        id="c1", document_id="doc1", section_id="s1", paragraph_id="p1",
        source_type="fact", source_id="f1", reference="Fait n°1", excerpt="excerpt",
    )
    return Document(
        id="doc1", template_id="consultation:v1", document_type=DocumentType.CONSULTATION,
        case_id=None, title="Consultation Test", sections=[section], citations=[citation],
    )


def test_export_returns_docx_media_type() -> None:
    result = DocxExporter().export(_document())
    assert result.format == ExportFormat.DOCX
    assert "wordprocessingml" in result.media_type
    assert result.filename == "doc1.docx"


def test_export_is_a_valid_docx_file_with_expected_text() -> None:
    result = DocxExporter().export(_document())
    document = docx.Document(io.BytesIO(result.content))
    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "Consultation Test" in all_text
    assert "Faits" in all_text
    assert "Le contrat a été rompu." in all_text
    assert "Fait n" in all_text
