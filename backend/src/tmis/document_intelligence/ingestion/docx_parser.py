import io

import docx

from tmis.document_intelligence.schemas.document import IngestedDocument


class DocxParser:
    """Implements `DocumentParserPort` for Word documents using
    `python-docx`. Paragraphs are joined with newlines so downstream
    layout analysis can work line-by-line, consistent with the other
    parsers."""

    content_types: tuple[str, ...] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def supports(self, content_type: str) -> bool:
        return content_type in self.content_types

    def parse(self, document_id: str, filename: str, raw_bytes: bytes) -> IngestedDocument:
        document = docx.Document(io.BytesIO(raw_bytes))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        core_properties = document.core_properties
        metadata = {}
        if core_properties.author:
            metadata["author"] = core_properties.author
        if core_properties.created:
            metadata["created_at"] = core_properties.created.isoformat()

        return IngestedDocument(
            id=document_id,
            filename=filename,
            content_type=self.content_types[0],
            text=text,
            page_count=1,
            raw_bytes=raw_bytes,
            metadata=metadata,
        )
