import io

import docx

from tmis.legal_drafting.citations.formatters import FootnoteCitationFormatter
from tmis.legal_drafting.documents.schemas import Document
from tmis.legal_drafting.export.schemas import ExportFormat, ExportResult

_DISCLAIMER = (
    "Document généré par TMIS — brouillon non validé, à relire et valider par un avocat."
)


class DocxExporter:
    """Implements `ExporterPort` via `python-docx`: preserves section
    structure as headings and attaches every citation right after the
    paragraph it supports (see docs/32-guide-exports.md)."""

    def export(self, document: Document) -> ExportResult:
        formatter = FootnoteCitationFormatter()
        citations_by_paragraph: dict[str, list[str]] = {}
        for citation in document.citations:
            citations_by_paragraph.setdefault(citation.paragraph_id, []).append(
                formatter.format(citation)
            )

        docx_document = docx.Document()
        disclaimer = docx_document.add_paragraph()
        disclaimer.add_run(_DISCLAIMER).italic = True
        docx_document.add_heading(document.title, level=1)

        for section in document.sections:
            docx_document.add_heading(section.title, level=2)
            for paragraph in section.paragraphs:
                docx_document.add_paragraph(paragraph.text)
                for formatted in citations_by_paragraph.get(paragraph.id, []):
                    citation_paragraph = docx_document.add_paragraph()
                    citation_paragraph.add_run(formatted).italic = True

        buffer = io.BytesIO()
        docx_document.save(buffer)
        return ExportResult(
            format=ExportFormat.DOCX,
            filename=f"{document.id}.docx",
            content=buffer.getvalue(),
            media_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )
